"""
Generate sample PDF and DOCX files for the DocuWise demo.
Run from the repo root: python scripts/make_samples.py
Requires: python-docx, fpdf2 (pip install python-docx fpdf2)
"""
import sys
from pathlib import Path

SAMPLES = Path(__file__).parent.parent / "samples"
SAMPLES.mkdir(exist_ok=True)


def make_docx():
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("python-docx not installed — skipping DOCX. pip install python-docx")
        return

    doc = Document()
    doc.add_heading("TechCorp Product Overview", 0)

    sections = [
        ("Executive Summary",
         "TechCorp's flagship product, DataPilot, is an AI-powered analytics platform that helps "
         "enterprises extract insights from structured and unstructured data. DataPilot processes "
         "over 2 billion events per month across 1,800+ customers worldwide."),
        ("Core Features",
         "DataPilot ships with real-time dashboards, anomaly detection, natural language querying, "
         "and automated report generation. The platform integrates with 120+ data sources including "
         "Snowflake, BigQuery, Salesforce, and all major cloud storage providers."),
        ("Pricing",
         "Starter: $499/month (up to 5 users, 10 data sources). "
         "Growth: $1,499/month (up to 25 users, unlimited sources). "
         "Enterprise: custom pricing, dedicated SLA, on-premise option available."),
        ("Security & Compliance",
         "DataPilot is SOC 2 Type II certified, GDPR compliant, and HIPAA eligible. "
         "All data is encrypted in transit (TLS 1.3) and at rest (AES-256). "
         "Single sign-on (SSO) supported via SAML 2.0 and OIDC."),
        ("Roadmap — Next 12 Months",
         "Q4 2024: Multi-model LLM support (GPT-4, Claude, Llama). "
         "Q1 2025: Mobile app (iOS + Android). "
         "Q2 2025: Collaborative dashboards with commenting. "
         "Q3 2025: On-premise deployment via Kubernetes Helm chart."),
        ("Support",
         "All plans include email support with 24-hour response SLA. "
         "Growth and Enterprise plans include live chat. "
         "Enterprise plans include a dedicated Customer Success Manager "
         "and quarterly business reviews."),
    ]

    for title, body in sections:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(body)
        p.style.font.size = Pt(11)
        doc.add_paragraph()

    out = SAMPLES / "product_overview.docx"
    doc.save(out)
    print(f"Created: {out}")


def make_pdf():
    try:
        from fpdf import FPDF
    except ImportError:
        print("fpdf2 not installed — skipping PDF. pip install fpdf2")
        return

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    W = pdf.w - pdf.l_margin - pdf.r_margin

    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(W, 12, "TechCorp Engineering Handbook", align="C")
    pdf.ln(16)

    sections = [
        ("Engineering Principles",
         ["Write code that a tired engineer can understand at 2am.",
          "Every production change requires a PR, two approvals, and a rollback plan.",
          "Incidents are blameless. We fix systems, not people.",
          "Prefer boring technology for infrastructure.",
          "Observability is not optional: every service must emit traces, metrics, and logs."]),
        ("On-Call Rotations",
         ["Each squad has a weekly primary and secondary on-call.",
          "Pages outside 9am-9pm local time are considered P1.",
          "On-call handoff happens every Monday at 10am UTC.",
          "Runbooks for every alert are mandatory.",
          "Post-mortems are due within 48 hours of incident resolution."]),
        ("Development Workflow",
         ["All work tracked in Linear. No ticket, no branch.",
          "Branch naming: type/ticket-id-short-desc.",
          "CI must pass before merge. No bypasses.",
          "Deploy to staging via PR merge. Deploy to prod manually.",
          "Feature flags for anything touching more than 10 percent of users."]),
        ("Infrastructure",
         ["Everything runs on AWS (eu-west-1 primary, us-east-1 failover).",
          "Kubernetes (EKS) for services, RDS Postgres for primary data.",
          "Terraform for all infra. No ClickOps.",
          "Secrets managed in AWS Secrets Manager. Never in committed env files.",
          "Monthly DR drill on the last Tuesday of each month."]),
        ("Code Review Standards",
         ["Reviews should complete within 1 business day.",
          "Comment tone: suggest, do not demand. Prefix nitpicks with nit:",
          "Performance implications must be called out for any hot path.",
          "Security review required for auth changes and data export.",
          "Architecture Decision Records for anything that will outlive the sprint."]),
    ]

    for title, bullets in sections:
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(W, 10, title)
        pdf.ln(12)
        pdf.set_font("Helvetica", "", 11)
        for bullet in bullets:
            pdf.multi_cell(W, 7, f"  - {bullet}")
        pdf.ln(4)

    out = SAMPLES / "engineering_handbook.pdf"
    pdf.output(str(out))
    print(f"Created: {out}")


if __name__ == "__main__":
    make_docx()
    make_pdf()
    print("Done. Files in samples/")
